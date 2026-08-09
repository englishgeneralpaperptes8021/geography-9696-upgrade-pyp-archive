# ********** 9696 Geography PYP Portal (8-Folder Architecture) 9/8/26 ***********
import datetime
import io
import os
import fitz  # PyMuPDF
import streamlit as st

# Word Document Libraries
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Google API Libraries
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload

# ==========================================
# 0. STREAMLIT PAGE CONFIG & CUSTOM STYLING
# ==========================================
st.set_page_config(
    page_title="9696 Geography PYP Portal", 
    page_icon="🌍",
    layout="wide"
)

# Custom CSS Theme (Earthy Geography Theme)
st.markdown("""
    <style>
    .stApp {
        background-color: #F4F6F0 !important;
    }
    [data-testid="stSidebar"] {
        background-color: #E2E8D8 !important;
    }
    [data-testid="stSidebar"] > div:first-child {
        background-color: #E2E8D8 !important;
    }
    div[data-baseweb="input"], 
    div[data-baseweb="select"] > div, 
    .stTextInput input, 
    .stSelectbox select {
        background-color: #FFFFFF !important;
        color: #1E4620 !important;
        border-radius: 8px !important;
        border: 1px solid #2E6F40 !important;
    }
    label, .stWidgetLabel p, h1, h2, h3, h4, h5, h6, p, span {
        color: #1E4620 !important;
        font-weight: bold !important;
    }
    input {
        color: #1E4620 !important;
    }
    .stButton button, .stDownloadButton button, [data-testid="baseButton-secondary"], [data-testid="baseButton-primary"] {
        background-color: #2E6F40 !important;
        color: #FFFFFF !important;
        border: 1px solid #1E4620 !important;
        font-weight: bold !important;
    }
    .stButton button:hover, .stDownloadButton button:hover {
        background-color: #1E4620 !important;
        color: #FFFFFF !important;
    }
    </style>
""", unsafe_allow_html=True)


# ==========================================
# 1. CONFIGURATION & DIRECTORY SETUP (8 FOLDERS)
# ==========================================
SYLLABUS_CODE = "9696"
SCOPES = ['https://www.googleapis.com/auth/drive.readonly']

LOCAL_FOLDERS = {
    "p1_physical": "9696_Paper1",
    "p2_human": "9696_Paper2",
    "p3_adv_physical": "9696_Paper3",
    "p4_adv_human": "9696_Paper4",
    "insert_p1_p3": "9696_InsertP1P3",
    "insert_p2_p4": "9696_InsertP2P4",
    "ms_p1_p2": "9696_MarkSchemeP1P2",
    "ms_p3_p4": "9696_MarkSchemeP3P4"
}

for folder_path in LOCAL_FOLDERS.values():
    os.makedirs(folder_path, exist_ok=True)

# ==========================================
# 2. SERVICE ACCOUNT AUTHENTICATION & SYNC
# ==========================================
def build_drive_service(write_access=False):
    """Authenticates using Google Service Account credentials stored in st.secrets."""
    try:
        if "gcp_service_account" in st.secrets:
            service_account_info = dict(st.secrets["gcp_service_account"])
            scopes = ['https://www.googleapis.com/auth/drive.file'] if write_access else SCOPES
            creds = service_account.Credentials.from_service_account_info(
                service_account_info, 
                scopes=scopes
            )
            return build('drive', 'v3', credentials=creds)
        else:
            st.error("❌ Missing [gcp_service_account] section in secrets.toml.")
            return None
    except Exception as e:
        st.error(f"❌ Service Account Authentication Error: {e}")
        return None

def sync_drive_folder_to_local(folder_key: str) -> tuple[int, str]:
    """Downloads missing files from Google Drive to local directories."""
    service = build_drive_service(write_access=False)
    if not service:
        return 0, "Failed to authenticate Service Account."
    
    folder_ids = st.secrets.get("drive_folders", {})
    drive_folder_id = folder_ids.get(folder_key)
    
    if not drive_folder_id:
        return 0, f"Missing drive_folder_id for `{folder_key}` in secrets.toml."

    local_path = LOCAL_FOLDERS[folder_key]
    
    try:
        query = f"'{drive_folder_id}' in parents and trashed = false"
        drive_files = []
        page_token = None

        while True:
            response = service.files().list(
                q=query,
                fields="nextPageToken, files(id, name, mimeType)",
                pageToken=page_token,
                pageSize=100
            ).execute()
            
            drive_files.extend(response.get('files', []))
            page_token = response.get('nextPageToken', None)
            
            if not page_token:
                break

        downloaded_count = 0

        for file_info in drive_files:
            file_name = file_info['name']
            file_id = file_info['id']
            local_file_path = os.path.join(local_path, file_name)

            if not os.path.exists(local_file_path):
                request = service.files().get_media(fileId=file_id)
                with open(local_file_path, "wb") as f:
                    downloader = MediaIoBaseDownload(f, request)
                    done = False
                    while not done:
                        _, done = downloader.next_chunk()
                downloaded_count += 1

        total_local_files = len([f for f in os.listdir(local_path) if os.path.isfile(os.path.join(local_path, f))])
        return downloaded_count, f"Synced {downloaded_count} new file(s) for `{folder_key}` (Total: {total_local_files})."
        
    except Exception as e:
        return 0, f"Sync error for `{folder_key}`: {e}"

def perform_bulk_sync():
    """Syncs all 8 configured Google Drive folders."""
    total_synced = 0
    messages = []
    for f_key in LOCAL_FOLDERS.keys():
        count, msg = sync_drive_folder_to_local(f_key)
        total_synced += count
        messages.append(msg)
    return total_synced, messages

# ==========================================
# 3. HELPER FUNCTIONS
# ==========================================
def create_worksheet_docx(basket_items: list) -> io.BytesIO:
    """Generates a Word document containing selected PDF pages."""
    doc = Document()
    section = doc.sections[0]

    section.page_width = Inches(8.5)
    section.page_height = Inches(11.5)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run("Page ")
    add_page_number_to_run(header_run)

    doc.add_heading(f'PTES {SYLLABUS_CODE} Geography Worksheet', level=1)

    for idx, item in enumerate(basket_items):
        doc.add_heading(f"Source: {item['file']} (Page {item['page'] + 1})", level=2)
        pdf_doc = fitz.open(item['path'])
        page = pdf_doc.load_page(item['page'])
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img_data = io.BytesIO(pix.tobytes("png"))

        doc.add_picture(img_data, width=Inches(7.8), height=Inches(8.8))

        if idx < len(basket_items) - 1:
            doc.add_page_break()
        pdf_doc.close()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def add_page_number_to_run(run):
    """Adds a dynamic Word page number field."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def render_pdf_page_preview(filepath: str, page_num: int = 0):
    """Renders a PDF page to PNG for in-app preview."""
    try:
        doc = fitz.open(filepath)
        page = doc.load_page(page_num)
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
        img_bytes = pix.tobytes("png")
        doc.close()
        return img_bytes
    except Exception as e:
        st.error(f"Unable to render page preview: {e}")
        return None

def execute_pdf_search(folder_key: str, keyword_string: str) -> list[dict]:
    """Searches PDF files in a specific folder for matching keywords."""
    results = []
    keywords = [k.strip().lower() for k in keyword_string.split(",") if k.strip()]
    folder_path = LOCAL_FOLDERS[folder_key]
    
    if os.path.exists(folder_path):
        for file in os.listdir(folder_path):
            if file.endswith(".pdf"):
                filepath = os.path.join(folder_path, file)
                try:
                    doc = fitz.open(filepath)
                    for page_num in range(len(doc)):
                        text = doc[page_num].get_text().lower()
                        if all(kw in text for kw in keywords):
                            results.append({
                                "file": file, 
                                "page": page_num, 
                                "path": filepath
                            })
                    doc.close()
                except Exception:
                    continue
    return results

# ==========================================
# 4. APP STATE INITIALIZATION
# ==========================================
if 'handout_basket' not in st.session_state:
    st.session_state.handout_basket = []

if 'p1_p3_results' not in st.session_state:
    st.session_state.p1_p3_results = []
if 'p2_p4_results' not in st.session_state:
    st.session_state.p2_p4_results = []
if 'insert_results' not in st.session_state:
    st.session_state.insert_results = []

if 'has_auto_synced' not in st.session_state:
    st.session_state.has_auto_synced = True
    with st.spinner("🚀 Waking up portal & auto-syncing Geography files via Service Account..."):
        perform_bulk_sync()

# ==========================================
# 5. STREAMLIT UI LAYOUT
# ==========================================
st.title("PUSAT TINGKATAN ENAM SENGKURONG")
st.subheader(f"🌍 {SYLLABUS_CODE} Geography PYP Resource Library")

# --- SIDEBAR CONTROLS ---
with st.sidebar:
    st.header("🔄 Google Drive Sync")
    if st.button("🔄 Sync Google Drive", type="primary", use_container_width=True):
        with st.spinner("Syncing Google Drive folders..."):
            synced_count, sync_msgs = perform_bulk_sync()
            st.success(f"🎉 Sync Complete! {synced_count} new file(s) downloaded.")
            for m in sync_msgs:
                st.caption(m)

    st.markdown("---")
    st.metric(label="Saved Pages in Basket", value=len(st.session_state.handout_basket))

    if st.button("🗑️ Clear Entire Basket", use_container_width=True):
        st.session_state.handout_basket = []
        st.rerun()

# --- NAVIGATION TABS ---
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "🏔️ Physical Geography (P1/P3)", 
    "🏙️ Human Geography (P2/P4)", 
    "🗺️ Inserts (P1-P4)",
    "🔑 Download Answer", 
    "🛒 Download Worksheet", 
    "⚙️ Upload PYP / Admin"
])

# --- TAB 1: PHYSICAL GEOGRAPHY (PAPER 1 or PAPER 3) ---
with tab1:
    st.subheader("🏔️ Physical Geography Search (Paper 1 / Paper 3)")
    
    selected_phys_paper = st.selectbox(
        "Select Component Paper:", 
        options=["p1_physical", "p3_adv_physical"],
        format_func=lambda x: "Paper 1 (12): Core Physical Geography" if x == "p1_physical" else "Paper 3 (32): Advanced Physical Geography Options",
        key="select_phys_paper"
    )

    t_kw = st.text_input(
        "Enter Keywords", 
        placeholder="e.g., hydrology, fluvial, atmosphere, weathering, coastal, hazardous", 
        key="phys_kw"
    )

    if st.button("Search Keyword", type="primary", key="btn_search_phys"):
        if t_kw.strip():
            with st.spinner("Scanning PDFs..."):
                st.session_state.p1_p3_results = execute_pdf_search(selected_phys_paper, t_kw)
        else:
            st.warning("Please enter a keyword.")

    if st.session_state.p1_p3_results:
        st.write(f"Found **{len(st.session_state.p1_p3_results)}** matching page(s):")
        for idx, item in enumerate(st.session_state.p1_p3_results):
            with st.expander(f"📄 {item['file']} | Page {item['page'] + 1}"):
                c1, c2 = st.columns([3, 1])
                with c1:
                    preview_img = render_pdf_page_preview(item["path"], item["page"])
                    if preview_img:
                        st.image(preview_img, use_container_width=True)
                with c2:
                    if st.button("➕ Add to Cart", key=f"add_phys_{idx}"):
                        st.session_state.handout_basket.append(item)
                        st.toast(f"Added Page {item['page'] + 1} to basket!")
                        st.rerun()
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    with open(item["path"], "rb") as pdf_f:
                        st.download_button(
                            label="📥 Download Full PDF",
                            data=pdf_f,
                            file_name=item["file"],
                            mime="application/pdf",
                            key=f"dl_phys_{idx}"
                        )

# --- TAB 2: HUMAN GEOGRAPHY (PAPER 2 or PAPER 4) ---
with tab2:
    st.subheader("🏙️ Human Geography Search (Paper 2 / Paper 4)")
    
    selected_human_paper = st.selectbox(
        "Select Component Paper:", 
        options=["p2_human", "p4_adv_human"],
        format_func=lambda x: "Paper 2 (22): Core Human Geography" if x == "p2_human" else "Paper 4 (42): Advanced Human Geography Options",
        key="select_human_paper"
    )

    h_kw = st.text_input(
        "Enter Keywords", 
        placeholder="e.g., population, migration, urban, settlement, energy, governance", 
        key="human_kw"
    )

    if st.button("Search Keyword", type="primary", key="btn_search_human"):
        if h_kw.strip():
            with st.spinner("Scanning PDFs..."):
                st.session_state.p2_p4_results = execute_pdf_search(selected_human_paper, h_kw)
        else:
            st.warning("Please enter a keyword.")

    if st.session_state.p2_p4_results:
        st.write(f"Found **{len(st.session_state.p2_p4_results)}** matching page(s):")
        for idx, item in enumerate(st.session_state.p2_p4_results):
            with st.expander(f"📄 {item['file']} | Page {item['page'] + 1}"):
                c1, c2 = st.columns([3, 1])
                with c1:
                    preview_img = render_pdf_page_preview(item["path"], item["page"])
                    if preview_img:
                        st.image(preview_img, use_container_width=True)
                with c2:
                    if st.button("➕ Add to Cart", key=f"add_human_{idx}"):
                        st.session_state.handout_basket.append(item)
                        st.toast(f"Added Page {item['page'] + 1} to basket!")
                        st.rerun()
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    with open(item["path"], "rb") as pdf_f:
                        st.download_button(
                            label="📥 Download Full PDF",
                            data=pdf_f,
                            file_name=item["file"],
                            mime="application/pdf",
                            key=f"dl_human_{idx}"
                        )

# --- TAB 3: INSERTS SEARCH (P1, P2, P3, P4) ---
with tab3:
    st.subheader("🗺️ Geography Inserts Search (Figures, Maps & Tables)")
    
    selected_insert = st.selectbox(
        "Select Insert Paper target:", 
        options=["insert_p1", "insert_p2", "insert_p3", "insert_p4"],
        format_func=lambda x: {
            "insert_p1": "Insert Paper 1 (Physical Core)",
            "insert_p2": "Insert Paper 2 (Human Core)",
            "insert_p3": "Insert Paper 3 (Advanced Physical)",
            "insert_p4": "Insert Paper 4 (Advanced Human)"
        }[x],
        key="select_insert_target"
    )

    ins_kw = st.text_input(
        "Enter Keywords or Fig/Table number", 
        placeholder="e.g., Fig 1.1, Table 2, hydrograph, satellite image", 
        key="insert_kw"
    )

    # Map selected insert option to its folder key
    folder_mapping = {
        "insert_p1": "insert_p1_p3",
        "insert_p3": "insert_p1_p3",
        "insert_p2": "insert_p2_p4",
        "insert_p4": "insert_p2_p4"
    }
    target_folder_key = folder_mapping[selected_insert]

    if st.button("Search Insert", type="primary", key="btn_search_insert"):
        if ins_kw.strip():
            with st.spinner("Scanning Inserts..."):
                st.session_state.insert_results = execute_pdf_search(target_folder_key, ins_kw)
        else:
            st.warning("Please enter a search keyword.")

    if st.session_state.insert_results:
        st.write(f"Found **{len(st.session_state.insert_results)}** matching insert page(s):")
        for idx, item in enumerate(st.session_state.insert_results):
            with st.expander(f"📄 {item['file']} | Page {item['page'] + 1}"):
                c1, c2 = st.columns([3, 1])
                with c1:
                    preview_img = render_pdf_page_preview(item["path"], item["page"])
                    if preview_img:
                        st.image(preview_img, use_container_width=True)
                with c2:
                    if st.button("➕ Add to Cart", key=f"add_ins_{idx}"):
                        st.session_state.handout_basket.append(item)
                        st.toast(f"Added Page {item['page'] + 1} to basket!")
                        st.rerun()
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    with open(item["path"], "rb") as pdf_f:
                        st.download_button(
                            label="📥 Download Full Insert PDF",
                            data=pdf_f,
                            file_name=item["file"],
                            mime="application/pdf",
                            key=f"dl_ins_{idx}"
                        )

# --- TAB 4: DOWNLOAD MARK SCHEME ---
with tab4:
    st.subheader("🔑 Download Answer / Mark Schemes")
    
    col_y, col_m, col_v = st.columns([1, 2, 2])
    with col_y:
        as_year = st.text_input(
            "Year", 
            value=str(datetime.datetime.now().year), 
            placeholder="YYYY", 
            key="as_pyp_year"
        )
    with col_m:
        as_month = st.selectbox("Select Session", ["June (s)", "November (w)"], key="as_mth")
        month_code = "s" if "June" in as_month else "w"
            
    with col_v:
        as_variant = st.selectbox(
            "Select Component Variant", 
            ["12", "22", "32", "42"], 
            key="as_var"
        )

    short_year = as_year.strip()[-2:] if len(as_year.strip()) >= 2 else ""
    expected_ms_filename = f"{SYLLABUS_CODE}_{month_code}{short_year}_ms_{as_variant}.pdf"

    st.markdown("---")
    found_ms_path = None
    
    # Search across both Mark Scheme folders
    for folder_key in ["ms_p1_p2", "ms_p3_p4"]:
        check_path = os.path.join(LOCAL_FOLDERS[folder_key], expected_ms_filename)
        if os.path.exists(check_path):
            found_ms_path = check_path
            break

    if found_ms_path:
        st.success(f"✅ Found Answer Scheme: `{expected_ms_filename}`")
        with open(found_ms_path, "rb") as f:
            st.download_button("📥 Download Answer Scheme PDF", f, file_name=expected_ms_filename, mime="application/pdf", type="primary")
        
        with st.expander("👁️ Preview Answer Scheme Document"):
            doc = fitz.open(found_ms_path)
            for p in range(len(doc)):
                img_data = render_pdf_page_preview(found_ms_path, p)
                if img_data:
                    st.image(img_data, caption=f"Page {p + 1}", use_container_width=True)
            doc.close()
    else:
        st.warning(f"⚠️ Mark Scheme `{expected_ms_filename}` was not found locally.")

# --- TAB 5: DOWNLOAD HANDOUT MERGED (CART) ---
with tab5:
    st.subheader("🛒 Download Handout Merged as Worksheet Management")
    
    if st.session_state.handout_basket:
        st.subheader("Selected Pages in Your Cart")
        st.markdown("Review your items below. Click **DELETE** to remove an individual page.")
        
        for idx, item in enumerate(st.session_state.handout_basket):
            col_info, col_action = st.columns([4, 1])
            with col_info:
                st.markdown(f"📄 **Item {idx + 1}:** `{item['file']}` — **Page {item['page'] + 1}**")
            with col_action:
                if st.button("🔴 DELETE", key=f"del_item_{idx}"):
                    st.session_state.handout_basket.pop(idx)
                    st.toast(f"Removed item {idx + 1} from cart.")
                    st.rerun()
            st.markdown("---")

        st.markdown("<br>", unsafe_allow_html=True)
        st.subheader("📝 Export Options")
        
        doc_buffer = create_worksheet_docx(st.session_state.handout_basket)
        target_filename = f"{SYLLABUS_CODE}_Geography_Worksheet.docx"

        st.download_button(
            label="🪄 Download Merged Word Document Worksheet",
            data=doc_buffer,
            file_name=target_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )
    else:
        st.info("🛒 Your cart is currently empty. Search for questions in Tabs 1–3 and click '➕ Add to Cart' to merge pages here.")

# --- TAB 6: UPLOAD PYP / ADMIN (8 FOLDERS) ---
with tab6:
    st.subheader("⚙️ Upload PYP / Admin Dashboard")
    st.caption("Secure admin controls for managing the 8 Google Drive repositories inside parent folder `9696GeogPYP`.")

    admin_pwd = st.secrets.get("ADMIN_PASSWORD", "")
    pwd_input = st.text_input("Enter Admin Password", type="password", key="admin_pwd_input")

    if pwd_input and pwd_input == admin_pwd:
        st.success("🔓 Authenticated as Administrator")
        st.markdown("---")
        
        admin_tab_links, admin_tab_upload = st.tabs(["📁 Drive Folders & Links", "📤 Direct File Upload"])
        
        drive_links = st.secrets.get("drive_web_links", {})
        drive_folder_ids = st.secrets.get("drive_folders", {})

        with admin_tab_links:
            st.markdown("### 🌐 Google Drive Web Dashboards")
            st.info("Click below to open any of your 8 Google Drive folders directly in your browser.")
            
            c1, c2, c3, c4 = st.columns(4)
            with c1:
                st.link_button("📘 Paper 1 Folder", drive_links.get("p1_physical", "https://drive.google.com"), use_container_width=True)
                st.link_button("🗺️ Insert P1/P3 Folder", drive_links.get("insert_p1_p3", "https://drive.google.com"), use_container_width=True)
            with c2:
                st.link_button("📗 Paper 2 Folder", drive_links.get("p2_human", "https://drive.google.com"), use_container_width=True)
                st.link_button("🗺️ Insert P2/P4 Folder", drive_links.get("insert_p2_p4", "https://drive.google.com"), use_container_width=True)
            with c3:
                st.link_button("📙 Paper 3 Folder", drive_links.get("p3_adv_physical", "https://drive.google.com"), use_container_width=True)
                st.link_button("🔑 MS P1/P2 Folder", drive_links.get("ms_p1_p2", "https://drive.google.com"), use_container_width=True)
            with c4:
                st.link_button("📕 Paper 4 Folder", drive_links.get("p4_adv_human", "https://drive.google.com"), use_container_width=True)
                st.link_button("🔑 MS P3/P4 Folder", drive_links.get("ms_p3_p4", "https://drive.google.com"), use_container_width=True)

        with admin_tab_upload:
            st.markdown("### ☁️ Direct Cloud Upload via Service Account")
            
            target_category = st.selectbox(
                "Select Destination Repository Folder", 
                options=list(LOCAL_FOLDERS.keys()),
                format_func=lambda x: {
                    "p1_physical": "9696_Paper1 (Core Physical)",
                    "p2_human": "9696_Paper2 (Core Human)",
                    "p3_adv_physical": "9696_Paper3 (Advanced Physical)",
                    "p4_adv_human": "9696_Paper4 (Advanced Human)",
                    "insert_p1_p3": "9696_InsertP1P3",
                    "insert_p2_p4": "9696_InsertP2P4",
                    "ms_p1_p2": "9696_MarkSchemeP1P2",
                    "ms_p3_p4": "9696_MarkSchemeP3P4"
                }[x],
                key="admin_upload_category"
            )
            
            uploaded_file = st.file_uploader("Choose a PDF file to upload", type=["pdf"], key="admin_file_uploader")
            
            if uploaded_file is not None:
                if st.button("🚀 Upload File to Google Drive", type="primary", key="execute_upload_btn"):
                    with st.spinner("Uploading file to Google Drive repository..."):
                        try:
                            service = build_drive_service(write_access=True)
                            if service:
                                target_folder_id = drive_folder_ids.get(target_category)
                                if not target_folder_id:
                                    st.error(f"❌ Missing drive folder ID for `{target_category}` in secrets.toml.")
                                else:
                                    file_metadata = {
                                        'name': uploaded_file.name,
                                        'parents': [target_folder_id]
                                    }
                                    
                                    temp_upload_path = os.path.join(LOCAL_FOLDERS[target_category], uploaded_file.name)
                                    with open(temp_upload_path, "wb") as f:
                                        f.write(uploaded_file.getbuffer())
                                        
                                    media = MediaFileUpload(temp_upload_path, mimetype='application/pdf', resumable=True)
                                    file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
                                    
                                    st.success(f"🎉 Successfully uploaded `{uploaded_file.name}` to Google Drive! (File ID: {file.get('id')})")
                        except Exception as e:
                            st.error(f"❌ Upload failed: {e}")

    elif pwd_input:
        st.error("❌ Incorrect Admin Password.")

# ==========================================
# 7. PORTAL FOOTER
# ==========================================
st.markdown("---")
SCHOOL_NAME = "Pusat Tingkatan Enam Sengkurong (PTES)"
SCHOOL_VISION = "Nurturing Resilient Leaders & Future-Ready Citizens"
DEVELOPER_NAME = "Cikgu Haziqah / Computer Science Department"

footer_html = f"""
<div style="text-align: center; padding: 15px 0px; font-family: sans-serif;">
    <p style="margin: 0; font-size: 1.0em; font-weight: bold;">🏫 {SCHOOL_NAME}</p>
    <p style="margin: 5px 0; font-size: 0.9em; font-style: italic;">"{SCHOOL_VISION}"</p>
    <p style="margin: 5px 0 0 0; font-size: 0.85em; font-weight: 600;">💻 Developed by {DEVELOPER_NAME}</p>
</div>
"""
st.markdown(footer_html, unsafe_allow_html=True)
